import os
import argparse
import logging
import re
import json

from pathlib import Path
from collections import defaultdict

import gspread
from google.auth import default

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

logger = logging.getLogger(__name__)


class GoogleSheets:
    SPREADSHEET_ID = "1Dnh4blW2lAa9PRYKVJL1i79JZCDc6V0Uzw2ZRR9P10Y"
    SCOPES = [
        "https://www.googleapis.com/auth/spreadsheets.readonly",
        "https://www.googleapis.com/auth/drive.readonly",
    ]

    @classmethod
    def google_authentication(cls):
        logger.info("Authenticating to Google Sheets . . .")

        if not os.getenv("GITHUB_ACTIONS"):
            # Running locally
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = str(Path("service-account-credentials.json").absolute())

        creds, _ = default(scopes=cls.SCOPES)
        return creds

    @classmethod
    def get_spreadsheet(cls, worksheet_name="Form Responses 1"):
        logger.info("Downloading spreadsheet . . .")

        creds = cls.google_authentication()
        gc = gspread.authorize(creds)
        spreadsheet = gc.open_by_key(cls.SPREADSHEET_ID).worksheet(worksheet_name)

        return spreadsheet

    @classmethod
    def process_spreadsheet_data(cls) -> dict:
        # NOTE: We're skipping most checks here as the data is guaranteed to be a table
        spreadsheet = cls.get_spreadsheet()
        spreadsheet_entries = spreadsheet.get_all_records()
        logger.info(f"Found {len(spreadsheet_entries)} records in the spreadsheet . . .")

        # NOTE: Mutate the data in the spreadsheet
        #       * Use the `numar_autorizatie` field as a primary key
        #       * Only record the most-recent entries
        #       * Log any duplicate entries
        translators_by_auth_no = {}
        for entry in spreadsheet_entries:
            # Gather fields
            auth_no = entry["Număr de autorizație"]
            phone_no = entry["Număr de telefon"]
            email = entry["Email Address"]
            other_contact_details = entry["Alte detalii"]
            DI = entry["Sunteți disponibil(ă) pentru interpretariat ? [DI]"]
            DII = entry["Sunteți disponibil(ă) pentru interpretariat la Instanțe ? [DII]"]
            LA = entry["Sunteți înscris pe lista vreunei Ambasade/vreunui Consulat [LA] ?"]

            # Mutate fields (as required)
            auth_no = remove_whitespace(auth_no)
            phone_no = clean_phone_number(phone_no)
            email = remove_whitespace(email).lower()
            DI = "DI" if DI == "Da" else ""
            DII = "DII" if DII == "Da" else ""
            LA = "LA" if LA == "Da" else ""
            other_contact_details = str(other_contact_details).strip()

            # Report duplicates as errors
            if auth_no in translators_by_auth_no.keys():
                logger.error(f"Duplicate translator found with Auth. No.: {auth_no}. Updating with latest info . . .")

            translators_by_auth_no[auth_no] = {
                "Nr. Aut.": auth_no,
                "Contact": "\n".join(filter(None, [phone_no, email, other_contact_details])),
                "DI/DII/LA": "\n".join(filter(None, [DI, DII, LA])),
            }

        return translators_by_auth_no


def process_mj_data(filtered_db_path: Path) -> dict:
    logger.info("Loading filtered MJ Database . . .")
    with open(filtered_db_path, "r") as f:
        mj_data = json.load(f)
    logger.info(f"Loaded {len(mj_data.keys())} records . . .")

    # NOTE: We're skipping most checks here as the data is guaranteed to be well-formed

    # Mutate fields
    logger.info("Processing filtered MJ Database . . .")
    mj_data_mutated = {}
    for auth_no, translator_details in mj_data.items():

        # Gather fields
        translator_name = translator_details["nume"]
        languages = translator_details["limbiAutorizate"]
        county = translator_details["judet"]
        appeals_court = translator_details["curteApel"]

        # Mutate field
        languages = sorted(languages.split(", "))
        county = (county or "").title()
        appeals_court = (appeals_court or "").title()

        mj_data_mutated[auth_no] = {
            "Nume": translator_name,
            "Limbă/Limbi": languages,
            "Județ/CA": "\n".join(filter(None, [county, appeals_court])),
        }

    return mj_data_mutated


def correlate_data(mj_data: dict, google_sheets_data: dict) -> dict:
    logger.info("Correlating filtered MJ Database with the Google Sheets data . . .")

    # Error log and ignore any entries that DO NOT appear in the MJ database
    invalid_auth_numbers = google_sheets_data.keys() - mj_data.keys()
    if len(invalid_auth_numbers) != 0:
        logger.error(f"The following Authorisation Numbers are invalid and are ignored: {invalid_auth_numbers}")
        for auth_no in invalid_auth_numbers:
            google_sheets_data.pop(auth_no)
    else:
        logger.info(f"All authorisation numbers are valid . . .")

    # Merge the two dicts by Authorisation Number
    correlated_data = {k: {**google_sheets_data[k], **mj_data[k]} for k in google_sheets_data.keys()}

    return correlated_data


def get_languages(data: dict) -> list:
    languages = [lang for v in data.values() for lang in v["Limbă/Limbi"]]
    languages = sorted(list(set(languages)))
    logger.info(f"Found {len(languages)} languages: {languages}")

    return languages


def group_data_by_language(data: dict) -> dict:
    logger.info(f"Group data by language . . .")
    languages = get_languages(data=data)

    # Group data by language
    translators_by_lang = defaultdict(list)
    for lang in languages:
        for translator in data.values():
            if lang in translator["Limbă/Limbi"]:
                translators_by_lang[lang].append(translator)

    # Handle special cases
    language_to_category_map = {
        # Dutch
        "Olandeză": "Olandeză/Neerlandeză",
        "Neerlandeză": "Olandeză/Neerlandeză",
        # Serbo-croatian
        "Sârbă": "Sârbă și croată",
        "Croată": "Sârbă și croată",
        "Sârbo-croată": "Sârbă și croată",
        # Hebrew
        "Ebraică": "Ebraică",
        "Ebraică(ivrit)": "Ebraică",
        # Greek
        "Greacă": "Greacă",
        "Neogreacă": "Greacă",
        "Greacă veche": "Greacă",
    }

    # Extract and remap special categories
    logger.info(f"Handle special languages . . .")
    category_map = defaultdict(list)
    for lang, lang_category in language_to_category_map.items():
        translators = translators_by_lang.pop(lang, None)
        if translators is not None:
            category_map[lang] += translators

    # Update original dict
    translators_by_lang.update(category_map)

    # Sort items by language, then by authorisation number (descendingly)
    logger.info(f"Sort dictionary by language and by authorisation number (descendingly) . . .")
    translators_by_lang = dict(sorted(translators_by_lang.items()))
    for translators in translators_by_lang.values():
        translators.sort(key=lambda x: int(x["Nr. Aut."]), reverse=True)

    return translators_by_lang


def generate_docx(data: dict):
    logger.info(f"Generating docx . . .")
    docx_document = Document()

    # Narrow margins (in cm)
    section = docx_document.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    # Set default font name and size
    style = docx_document.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(10)

    # Fill with data
    language_to_code = {
        "Albaneză": "SQ",
        "Arabă": "AR",
        "Armeană": "HY",
        "Bulgară": "BG",
        "Catalană": "CA",
        "Cehă": "CZ",
        "Chineză": "ZH",
        "Sârbo-croată": "SH",
        "Sârbă și croată": "SR/HR/SH",
        "Croată": "HR",
        "Sârbă": "SR",
        "Daneză": "DK",
        "Ebraică": "HE",
        "Ebraică(ivrit)": "HE",
        "Engleză": "EN",
        "Finlandeză": "FI",
        "Franceză": "FR",
        "Germană": "DE",
        "Greacă": "EL",
        "Neogreacă": "EL",
        "Greacă veche": "greaca veche",
        "Italiană": "IT",
        "Japoneză": "JA",
        "Latină": "LA",
        "Lituaniană": "LT",
        "Macedoneană": "MK",
        "Maghiară": "HU",
        "Neerlandeză": "NL",
        "Olandeză": "NL",
        "Olandeză/Neerlandeză": "NL",
        "Norvegiană": "NO",
        "Persană": "FA",
        "Polonă": "PL",
        "Portugheză": "PT",
        "Rromani": "RMN",
        "Rusă": "RU",
        "Slovacă": "SK",
        "Slovenă": "SL",
        "Spaniolă": "ES",
        "Suedeză": "SU",
        "Turcă": "TR",
        "Ucraineană": "UA",
    }

    # Heading and Legend
    heading = docx_document.add_heading("Lista traducătorilor activi", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(20)

    docx_document.add_heading("Legenda", level=2)
    docx_document.add_paragraph("DI = disponibil(ă) pentru interpretariat", style="List Bullet")
    docx_document.add_paragraph("DII = disponibil(ă) pentru interpretariat la Instanțe", style="List Bullet")
    docx_document.add_paragraph("LA = înscris(ă) pe lista la Ambasadă/Consulat", style="List Bullet")

    # Tables
    for lang, translators_list in data.items():
        # Add subtitle
        docx_document.add_heading(f"{lang} ({language_to_code[lang]})", level=2)

        # Add Table
        table = docx_document.add_table(
            rows=1 + len(translators_list), cols=len(list(translators_list[0].keys()))
        )  # NOTE: 1+ to account for the header too
        table.style = "Table Grid"

        # Add header
        header = table.rows[0].cells
        header[0].text = "Nr. Aut."
        header[1].text = "Nume"
        header[2].text = "Limbă/Limbi"
        header[3].text = "Județ/CA"
        header[4].text = "DI/DII/LA"
        header[5].text = "Contact"

        for idx, translator in enumerate(translators_list):
            cells = table.rows[idx + 1].cells  # NOTE: Offset from the header

            cells[0].text = translator["Nr. Aut."]
            cells[1].text = translator["Nume"]
            cells[2].text = ", ".join([language_to_code[x] for x in translator["Limbă/Limbi"]])
            cells[3].text = translator["Județ/CA"]
            cells[4].text = translator["DI/DII/LA"]
            cells[5].text = translator["Contact"]

        # Centre content in all cells
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return docx_document


def parse_args():
    logger.info("Parsing arguments . . .")

    parser = argparse.ArgumentParser()
    parser.add_argument("-f", "--filtered-db-path", type=Path, default=Path("output/translators_filtered.json"))
    parser.add_argument("-s", "--save-path", type=Path, default=Path("output/active_translators.docx"))
    args = parser.parse_args()

    return args


def remove_whitespace(text) -> str:
    return re.sub(r"\s+", "", str(text))


def clean_phone_number(phone_number) -> str:
    phone_number = remove_whitespace(phone_number)

    if phone_number.startswith("40"):  # The "+" from "+40" got escaped
        phone_number = f"+{phone_number}"

    return phone_number


def main():
    args = parse_args()

    # Correlate data between the MJ database and the Google Sheets worksheet
    mj_data = process_mj_data(args.filtered_db_path)
    google_sheets_data = GoogleSheets.process_spreadsheet_data()
    correlated_data = correlate_data(mj_data=mj_data, google_sheets_data=google_sheets_data)

    # Group data by language
    translators_by_lang = group_data_by_language(data=correlated_data)

    # Generate .docx file
    docx_document = generate_docx(data=translators_by_lang)

    # Save .docx file
    logger.info(f"Saving .docx file at {args.save_path}")
    docx_document.save(args.save_path)

    return


if __name__ == "__main__":
    main()
